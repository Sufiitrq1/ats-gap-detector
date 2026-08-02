"""
Extracts plain text from an uploaded resume file (PDF or DOCX).
"""
import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # also pull text out of any tables (skills tables etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs).strip()


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif lower.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.")

    if not text or len(text) < 50:
        raise ValueError(
            "Couldn't extract meaningful text from this resume. "
            "If it's a scanned/image PDF, try exporting a text-based version."
        )
    return text
